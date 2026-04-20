from __future__ import annotations

import os
from typing import Any
 
 
def _guess_title(*, filename: str, form_title: str | None) -> str:
    t = (form_title or "").strip()
    if t:
        return t
    base = os.path.splitext(filename)[0].strip()
    return base or "未命名文档"
 
 
def bytes_to_text(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("utf-8", errors="replace")
 
 
def _normalize_md(text: str) -> str:
    out = text.replace("\r\n", "\n").replace("\r", "\n")
    return out.strip()
 
 
def _text_to_md(text: str) -> str:
    return _normalize_md(text)
 
 
def docx_bytes_to_md(raw: bytes) -> str:
    try:
        from docx import Document as DocxDocument  # type: ignore[import-not-found]
    except Exception as e:
        raise RuntimeError("missing_dependency:python-docx") from e
 
    from io import BytesIO
 
    doc = DocxDocument(BytesIO(raw))
    lines: list[str] = []
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        style_name = (p.style.name or "") if p.style is not None else ""
        if style_name.lower().startswith("heading"):
            level = 1
            try:
                parts = style_name.split()
                if len(parts) >= 2:
                    level = int(parts[1])
            except Exception:
                level = 1
            level = max(1, min(6, level))
            lines.append("#" * level + " " + txt)
        else:
            lines.append(txt)
        lines.append("")
    return _normalize_md("\n".join(lines))
 
 
def pdf_bytes_to_md(raw: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except Exception as e:
        raise RuntimeError("missing_dependency:pypdf") from e
 
    from io import BytesIO
 
    reader = PdfReader(BytesIO(raw))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        txt = page.extract_text() or ""
        txt = txt.strip()
        if not txt:
            continue
        if i > 0:
            parts.append("")
        parts.append(txt)
    return _normalize_md("\n\n".join(parts))
 
 
def convert_upload_to_markdown(*, filename: str, raw: bytes) -> tuple[str, str]:
    ext = os.path.splitext(filename)[1].lower()
    if ext in {".md", ".markdown"}:
        return _normalize_md(bytes_to_text(raw)), "md"
    if ext == ".txt":
        return _text_to_md(bytes_to_text(raw)), "txt"
    if ext == ".docx":
        return docx_bytes_to_md(raw), "docx"
    if ext == ".pdf":
        return pdf_bytes_to_md(raw), "pdf"
    if ext == ".doc":
        raise ValueError("unsupported_doc")
    raise ValueError("unsupported_file")
 
 
def md_to_text(md: str) -> str:
    text = _normalize_md(md)
    lines: list[str] = []
    in_code = False
    for raw in text.split("\n"):
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            lines.append(line)
            continue
 
        s = line.lstrip()
        while s.startswith("#"):
            s = s[1:]
        s = s.lstrip()
 
        if s.startswith(("- ", "* ")):
            s = s[2:].lstrip()
 
        if s.startswith(">"):
            s = s[1:].lstrip()
 
        lines.append(s)
    return "\n".join(lines).strip() + "\n"
 
 
def md_to_docx_bytes(md: str, *, title: str | None = None) -> bytes:
    try:
        from docx import Document as DocxDocument  # type: ignore[import-not-found]
    except Exception as e:
        raise RuntimeError("missing_dependency:python-docx") from e
 
    from io import BytesIO
 
    doc = DocxDocument()
    if title and title.strip():
        doc.add_heading(title.strip(), level=1)
 
    text = _normalize_md(md)
    in_code = False
    for raw in text.split("\n"):
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
 
        if in_code:
            doc.add_paragraph(line)
            continue
 
        s = line.strip()
        if not s:
            doc.add_paragraph("")
            continue
 
        if s.startswith("#"):
            level = len(s) - len(s.lstrip("#"))
            level = max(1, min(6, level))
            heading = s[level:].strip()
            doc.add_heading(heading, level=level)
            continue
 
        if s.startswith(("- ", "* ")):
            p = doc.add_paragraph(s[2:].strip())
            try:
                p.style = "List Bullet"
            except Exception:
                pass
            continue
 
        doc.add_paragraph(s)
 
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
 
 
def md_to_pdf_bytes(md: str, *, title: str | None = None) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore[import-not-found]
        from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-not-found]
        from reportlab.pdfbase import pdfmetrics  # type: ignore[import-not-found]
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore[import-not-found]
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore[import-not-found]
    except Exception as e:
        raise RuntimeError("missing_dependency:reportlab") from e
 
    from io import BytesIO
    from xml.sax.saxutils import escape
 
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=(title or "").strip() or None)
 
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    base = styles["BodyText"]
    base.fontName = "STSong-Light"
    base.fontSize = 11
    base.leading = 16
 
    h1 = styles["Heading1"]
    h1.fontName = "STSong-Light"
    h2 = styles["Heading2"]
    h2.fontName = "STSong-Light"
 
    story: list[Any] = []
    if title and title.strip():
        story.append(Paragraph(escape(title.strip()), h1))
        story.append(Spacer(1, 10))
 
    text = _normalize_md(md)
    in_code = False
    for raw in text.split("\n"):
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if not line.strip():
            story.append(Spacer(1, 6))
            continue
 
        if not in_code and line.lstrip().startswith("#"):
            s = line.lstrip()
            level = len(s) - len(s.lstrip("#"))
            heading = s[level:].strip()
            story.append(Paragraph(escape(heading), h1 if level <= 1 else h2))
            story.append(Spacer(1, 6))
            continue
 
        content = md_to_text(line) if not in_code else line
        story.append(Paragraph(escape(content), base))
 
    doc.build(story)
    return buf.getvalue()
 
 
def safe_filename(name: str) -> str:
    s = (name or "").strip() or "document"
    out: list[str] = []
    for ch in s:
        if ch.isascii() and (ch.isalnum() or ch in {"-", "_"}):
            out.append(ch)
        elif ch.isspace():
            out.append("_")
    return ("".join(out) or "document")[:80]
